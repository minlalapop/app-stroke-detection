import html
import json
import uuid
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from sqlalchemy import select
from sqlalchemy.orm import Session
from weasyprint import HTML

from app.models.export import ExportFormat, ExportSource, ReportExport
from app.models.report import Report, ReportStatus
from app.models.user import User
from app.storage.file_storage import LocalFileStorage


class ReportNotFoundError(Exception):
    pass


class ReportNotReadyForExportError(Exception):
    pass


class ExportNotFoundError(Exception):
    pass


class ExportService:
    def __init__(self, db: Session, storage: LocalFileStorage) -> None:
        self.db = db
        self.storage = storage

    def create(
        self, report_id: uuid.UUID, export_format: ExportFormat, generated_by: User
    ) -> ReportExport:
        report = self.db.get(Report, report_id)
        if report is None:
            raise ReportNotFoundError
        if report.status != ReportStatus.READY_FOR_EXPORT:
            raise ReportNotReadyForExportError
        data = json.loads(self.storage.read(report.deterministic_report_uri))
        llm_text = None
        source = ExportSource.DETERMINISTIC
        if report.llm_enriched_report_uri and report.llm_approved_at:
            llm_text = self.storage.read(report.llm_enriched_report_uri).decode("utf-8")
            source = ExportSource.DETERMINISTIC_WITH_APPROVED_LLM

        content = self._render(export_format, data, llm_text)
        export_id = uuid.uuid4()
        stored = self.storage.save_export(
            report.id, export_id, content, f".{export_format.value}"
        )
        export = ReportExport(
            id=export_id,
            report_id=report.id,
            format=export_format,
            source=source,
            file_uri=stored.uri,
            sha256=stored.sha256,
            generated_by=generated_by.id,
        )
        try:
            self.db.add(export)
            self.db.commit()
            self.db.refresh(export)
        except Exception:
            self.db.rollback()
            self.storage.delete(stored.uri)
            raise
        return export

    def list_for_report(self, report_id: uuid.UUID) -> list[ReportExport]:
        if self.db.get(Report, report_id) is None:
            raise ReportNotFoundError
        return list(
            self.db.scalars(
                select(ReportExport)
                .where(ReportExport.report_id == report_id)
                .order_by(ReportExport.generated_at, ReportExport.id)
            ).all()
        )

    def get(self, export_id: uuid.UUID) -> ReportExport:
        export = self.db.get(ReportExport, export_id)
        if export is None:
            raise ExportNotFoundError
        return export

    def _render(self, export_format: ExportFormat, data: dict, llm_text: str | None) -> bytes:
        if export_format == ExportFormat.PDF:
            return self._pdf(data, llm_text)
        if export_format == ExportFormat.DOCX:
            return self._docx(data, llm_text)
        return self._xlsx(data, llm_text)

    def _pdf(self, data: dict, llm_text: str | None) -> bytes:
        sections = self._sections(data)
        section_html = "".join(
            f"<h2>{html.escape(title)}</h2><table>"
            + "".join(
                f"<tr><th>{html.escape(label)}</th><td>{html.escape(self._text(value))}</td></tr>"
                for label, value in rows
            )
            + "</table>"
            for title, rows in sections
        )
        llm_html = ""
        if llm_text:
            llm_html = (
                "<h2>Enrichissement LLM approuvé</h2>"
                f"<div class='llm'>{html.escape(llm_text)}</div>"
            )
        document = f"""
        <html><head><meta charset="utf-8"><style>
        @page {{ size: A4; margin: 18mm; }}
        body {{ font-family: sans-serif; color: #17212b; font-size: 10pt; }}
        h1 {{ color: #126782; font-size: 20pt; margin-bottom: 4mm; }}
        h2 {{ color: #126782; font-size: 13pt; margin-top: 6mm; margin-bottom: 2mm; }}
        table {{ width: 100%; border-collapse: collapse; page-break-inside: avoid; }}
        th, td {{ border: 1px solid #cbd5df; padding: 6px; vertical-align: top; }}
        th {{ width: 38%; background: #eaf5f8; text-align: left; }}
        .warning {{ background: #fff4cc; border-left: 4px solid #d99a00; padding: 8px; }}
        .llm {{ white-space: pre-wrap; border: 1px solid #8ab8c7; padding: 10px; }}
        </style></head><body>
        <h1>Rapport d’aide à l’analyse de l’AVC</h1>
        <p class="warning">Outil d’aide à la décision — ne remplace pas l’avis médical.</p>
        {section_html}{llm_html}
        </body></html>
        """
        return HTML(string=document).write_pdf()

    def _docx(self, data: dict, llm_text: str | None) -> bytes:
        document = Document()
        section = document.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        document.styles["Normal"].font.size = Pt(9)
        document.styles["Heading 1"].font.size = Pt(12)
        document.styles["Heading 1"].font.color.rgb = RGBColor(18, 103, 130)
        document.styles["Heading 1"].paragraph_format.space_before = Pt(6)
        document.styles["Heading 1"].paragraph_format.space_after = Pt(2)
        title = document.add_heading("Rapport d’aide à l’analyse de l’AVC", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.color.rgb = RGBColor(18, 103, 130)
        warning = document.add_paragraph()
        run = warning.add_run("Outil d’aide à la décision — ne remplace pas l’avis médical.")
        run.bold = True
        for section_title, rows in self._sections(data):
            document.add_heading(section_title, level=1)
            for label, value in rows:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(3)
                label_run = paragraph.add_run(f"{label} : ")
                label_run.bold = True
                paragraph.add_run(self._text(value))
        if llm_text:
            document.add_heading("Enrichissement LLM approuvé", level=1)
            document.add_paragraph(llm_text)
        footer = section.footer.paragraphs[0]
        footer.text = "Rapport généré par la plateforme d’aide à l’analyse de l’AVC"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.size = Pt(8)
        output = BytesIO()
        document.save(output)
        return output.getvalue()

    def _xlsx(self, data: dict, llm_text: str | None) -> bytes:
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Rapport"
        sheet.sheet_view.showGridLines = False
        sheet.freeze_panes = "A3"
        sheet.merge_cells("A1:B1")
        sheet["A1"] = "Rapport d’aide à l’analyse de l’AVC"
        sheet["A1"].font = Font(size=16, bold=True, color="FFFFFF")
        sheet["A1"].fill = PatternFill("solid", fgColor="126782")
        sheet["A1"].alignment = Alignment(horizontal="center")
        row = 3
        for section_title, rows in self._sections(data):
            sheet.merge_cells(start_row=row, start_column=1, end_row=row, end_column=2)
            cell = sheet.cell(row=row, column=1, value=section_title)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="3A8DA8")
            row += 1
            for label, value in rows:
                sheet.cell(row=row, column=1, value=label).font = Font(bold=True)
                sheet.cell(row=row, column=1).fill = PatternFill("solid", fgColor="EAF5F8")
                sheet.cell(row=row, column=2, value=self._text(value))
                row += 1
            row += 1
        sheet.column_dimensions["A"].width = 34
        sheet.column_dimensions["B"].width = 70
        for cells in sheet.iter_rows():
            for cell in cells:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        if llm_text:
            llm_sheet = workbook.create_sheet("Enrichissement LLM")
            llm_sheet["A1"] = "Enrichissement LLM approuvé"
            llm_sheet["A1"].font = Font(size=14, bold=True, color="FFFFFF")
            llm_sheet["A1"].fill = PatternFill("solid", fgColor="126782")
            llm_sheet["A3"] = llm_text
            llm_sheet["A3"].alignment = Alignment(wrap_text=True, vertical="top")
            llm_sheet.column_dimensions["A"].width = 100
        output = BytesIO()
        workbook.save(output)
        return output.getvalue()

    @staticmethod
    def _sections(data: dict) -> list[tuple[str, list[tuple[str, object]]]]:
        patient = data.get("patient") or {}
        clinical = data.get("clinical_data") or {}
        tabular = data.get("tabular_result") or {}
        imaging = data.get("imaging_result") or {}
        validation = data.get("doctor_validation") or {}
        return [
            ("Identification", [("Analyse", data.get("analysis_id")), ("Date", data.get("generated_at"))]),
            ("Patient", [("Identifiant", patient.get("id")), ("Prénom", patient.get("first_name")), ("Nom", patient.get("last_name")), ("Date de naissance", patient.get("birth_date")), ("Sexe", patient.get("sex"))]),
            ("Données cliniques", list(clinical.items()) or [("Statut", "Non disponibles")]),
            ("Résultat tabulaire", list(tabular.items()) or [("Statut", "Non disponible")]),
            ("Résultat image", list(imaging.items()) or [("Statut", "Non disponible")]),
            ("Validation médecin", [("Statut", validation.get("status")), ("Commentaire", validation.get("comment")), ("Médecin", validation.get("doctor_name")), ("Date", validation.get("validated_at"))]),
            ("Limitations", [("Information", item) for item in data.get("limitations", [])]),
        ]

    @staticmethod
    def _text(value: object) -> str:
        if value is None:
            return "Non disponible"
        if isinstance(value, bool):
            return "Oui" if value else "Non"
        return str(value)
